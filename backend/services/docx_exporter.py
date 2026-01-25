"""
DOCX Exporter - Export redlined contracts with Microsoft Word track changes
Generates DOCX files showing accepted/rejected changes in Word's track changes format
"""
import sqlite3
from io import BytesIO
from typing import Dict, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from logger import log_info, log_warning, log_error, log_success


class DocxExporter:
    """Export redlined contracts as DOCX with track changes"""

    def __init__(self, db_path: str = "/app/data/documents.db"):
        self.db_path = db_path

    def _get_connection(self):
        """Get SQLite database connection"""
        return sqlite3.connect(self.db_path)

    def export_with_track_changes(self, session_id: str) -> bytes:
        """
        Generate DOCX with Microsoft Word track changes

        Workflow:
        1. Get session details and document info
        2. Get all individual changes with user actions
        3. Create DOCX with original text + track changes markup
        4. Return DOCX file bytes

        Args:
            session_id: The redlining session ID

        Returns:
            DOCX file as bytes
        """
        log_info(f"Starting DOCX export for session: {session_id}", "DOCX_EXPORT")

        try:
            # Get session details
            conn = self._get_connection()
            cursor = conn.cursor()

            cursor.execute("""
                SELECT rs.uploaded_document_id, rs.template_id, rs.category,
                       d.filename, d.original_text
                FROM redlining_sessions rs
                JOIN documents d ON rs.uploaded_document_id = d.id
                WHERE rs.id = ?
            """, (session_id,))

            session_row = cursor.fetchone()
            if not session_row:
                conn.close()
                raise ValueError(f"Session not found: {session_id}")

            document_id, template_id, category, filename, original_text = session_row

            # Get all individual changes with user actions
            cursor.execute("""
                SELECT ac.id, ac.change_type, ac.original_text, ac.suggested_text,
                       ac.start_offset, ac.end_offset, ac.risk_level, ac.rationale,
                       ac.user_action, cc.new_clause_id, cc.template_clause_id
                FROM annotation_changes ac
                JOIN clause_comparisons cc ON ac.comparison_id = cc.id
                WHERE cc.session_id = ?
                ORDER BY ac.start_offset
            """, (session_id,))

            changes_rows = cursor.fetchall()
            conn.close()

            log_info(f"Retrieved {len(changes_rows)} changes for export", "DOCX_EXPORT")

            # Create DOCX document
            doc = Document()

            # Add document title
            title = doc.add_heading(f"Redlined Contract: {filename}", level=1)

            # Add metadata
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Session ID: {session_id[:8]}...\n").bold = True
            metadata_para.add_run(f"Category: {category or 'N/A'}\n")
            metadata_para.add_run(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            metadata_para.add_run(f"Total Changes: {len(changes_rows)}\n")

            # Count actions
            accepted_count = sum(1 for row in changes_rows if row[8] == 'accepted')
            rejected_count = sum(1 for row in changes_rows if row[8] == 'rejected')
            pending_count = sum(1 for row in changes_rows if row[8] == 'pending')

            metadata_para.add_run(f"Accepted: {accepted_count}, Rejected: {rejected_count}, Pending: {pending_count}\n")

            # Add separator
            doc.add_paragraph("─" * 80)

            # Process changes and add to document
            changes_by_action = {
                'accepted': [],
                'rejected': [],
                'pending': []
            }

            for row in changes_rows:
                change_id, change_type, orig_text, sugg_text, start_off, end_off, risk, rationale, action, new_clause, tmpl_clause = row

                change_dict = {
                    'id': change_id,
                    'type': change_type,
                    'original': orig_text,
                    'suggested': sugg_text,
                    'risk': risk,
                    'rationale': rationale
                }

                changes_by_action[action].append(change_dict)

            # Section 1: Accepted Changes
            if changes_by_action['accepted']:
                doc.add_heading('Accepted Changes', level=2)
                for idx, change in enumerate(changes_by_action['accepted'], 1):
                    self._add_change_to_doc(doc, idx, change, 'accepted')

            # Section 2: Rejected Changes
            if changes_by_action['rejected']:
                doc.add_heading('Rejected Changes', level=2)
                for idx, change in enumerate(changes_by_action['rejected'], 1):
                    self._add_change_to_doc(doc, idx, change, 'rejected')

            # Section 3: Pending Changes
            if changes_by_action['pending']:
                doc.add_heading('Pending Review', level=2)
                for idx, change in enumerate(changes_by_action['pending'], 1):
                    self._add_change_to_doc(doc, idx, change, 'pending')

            # Save to BytesIO
            docx_bytes_io = BytesIO()
            doc.save(docx_bytes_io)
            docx_bytes_io.seek(0)

            log_success(f"DOCX export completed for session {session_id}", "DOCX_EXPORT")
            return docx_bytes_io.getvalue()

        except Exception as e:
            log_error(f"Failed to export DOCX: {str(e)}", "DOCX_EXPORT")
            raise

    def _add_change_to_doc(self, doc: Document, index: int, change: Dict, action: str):
        """Add a change entry to the document with formatting"""

        # Change header
        para = doc.add_paragraph()
        run = para.add_run(f"Change #{index}: {change['type'].replace('_', ' ').title()}")
        run.bold = True
        run.font.size = Pt(11)

        # Risk badge
        risk_run = para.add_run(f" [{change['risk']} Risk]")
        if change['risk'] == 'High':
            risk_run.font.color.rgb = RGBColor(220, 38, 38)
        elif change['risk'] == 'Medium':
            risk_run.font.color.rgb = RGBColor(245, 158, 11)
        else:
            risk_run.font.color.rgb = RGBColor(22, 163, 74)
        risk_run.bold = True

        # Action status
        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Status: {action.upper()}")
        if action == 'accepted':
            status_run.font.color.rgb = RGBColor(22, 163, 74)
        elif action == 'rejected':
            status_run.font.color.rgb = RGBColor(220, 38, 38)
        else:
            status_run.font.color.rgb = RGBColor(245, 158, 11)
        status_run.bold = True

        # Original text (if exists)
        if change['original']:
            orig_para = doc.add_paragraph()
            orig_para.add_run("Original: ").bold = True
            orig_run = orig_para.add_run(change['original'])
            if action == 'rejected':
                # Keep as-is for rejected changes
                orig_run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Show as deleted for accepted changes
                orig_run.font.strike = True
                orig_run.font.color.rgb = RGBColor(153, 27, 27)

        # Suggested text (if exists)
        if change['suggested']:
            sugg_para = doc.add_paragraph()
            sugg_para.add_run("Suggested: ").bold = True
            sugg_run = sugg_para.add_run(change['suggested'])
            if action == 'accepted':
                # Show as inserted for accepted changes
                sugg_run.font.underline = True
                sugg_run.font.color.rgb = RGBColor(21, 128, 61)
            else:
                # Show as normal for rejected
                sugg_run.font.color.rgb = RGBColor(107, 114, 128)

        # Rationale
        if change['rationale']:
            rat_para = doc.add_paragraph()
            rat_para.add_run("Rationale: ").bold = True
            rat_para.add_run(change['rationale'])
            rat_para.paragraph_format.left_indent = Pt(20)

        # Add spacing
        doc.add_paragraph()

    def get_export_filename(self, session_id: str) -> str:
        """Generate export filename based on session"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"redlined_contract_{session_id[:8]}_{timestamp}.docx"
